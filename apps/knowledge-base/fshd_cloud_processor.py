import os
import sys
from pathlib import Path

current_file = Path(__file__).resolve()
project_root = current_file.parent.parent.parent  # knowledge-base -> apps -> openrd-master
sys.path.insert(0, str(project_root))

print(f"📁 项目根目录: {project_root}")
print(f"📁 当前目录: {current_file.parent}")

try:
    from config import config
    print(f"✅ 成功导入配置模块")
    
    # 验证配置
    config.chroma.validate()
    
except ImportError as e:
    print(f"❌ 导入配置模块失败: {e}")
    print("💡 请确保项目根目录有 config/__init__.py 文件")
    print(f"💡 Python路径: {sys.path}")
    sys.exit(1)
except ValueError as e:
    print(f"❌ 配置验证失败: {e}")
    sys.exit(1)

import json
import time
import chromadb
import PyPDF2
from docx import Document
from datetime import datetime
from typing import List, Dict
import re
import hashlib
import langdetect
from langdetect import detect, DetectorFactory

DetectorFactory.seed = 0

class FSHDBatchProcessor:
    """基础的云端FSHD知识库处理器"""
    def __init__(self, cloud_api_key: str, tenant_id: str, database_name: str = "FSHD"):
        """初始化云端FSHD知识库处理器"""
        self.client = chromadb.CloudClient(
            api_key=cloud_api_key,
            tenant=tenant_id,
            database=database_name
        )
        self.collection = self.client.get_or_create_collection("fshd_knowledge_base")
        print("🚀 云端FSHD知识库处理器初始化完成！")
        print(f"📍 连接至: {database_name} | 租户: {tenant_id}")
    
    def generate_short_id(self, category: str, filename: str, language: str, chunk_index: int) -> str:
        """生成短ID以避免配额限制"""
        base_string = f"{category}_{filename}_{language}_{chunk_index}"
        return hashlib.md5(base_string.encode()).hexdigest()[:16]
    
    def detect_language(self, text: str) -> str:
        """检测文本语言"""
        try:
            sample_text = text[:1000] if len(text) > 1000 else text
            if len(sample_text.strip()) < 10:  
                return "unknown"
            return detect(sample_text)
        except:
            return "unknown"
    
    def extract_text_from_pdf(self, pdf_path: str) -> Dict:
        """从PDF提取文本"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                total_pages = len(reader.pages)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
            return {
                "text": text,
                "success": True,
                "pages": total_pages,
                "file_type": "pdf"
            }
        except Exception as e:
            print(f"❌ 读取PDF失败 {pdf_path}: {e}")
            return {"success": False, "error": str(e)}
    
    def extract_text_from_docx(self, docx_path: str) -> Dict:
        """从Word文档提取文本"""
        text = ""
        try:
            doc = Document(docx_path)
            total_paragraphs = len(doc.paragraphs)
            for para in doc.paragraphs:
                text += para.text + "\n"
            return {
                "text": text,
                "success": True,
                "paragraphs": total_paragraphs,
                "file_type": "docx"
            }
        except Exception as e:
            print(f"❌ 读取Word文档失败 {docx_path}: {e}")
            return {"success": False, "error": str(e)}
    
    def extract_text(self, file_path: str) -> Dict:
        """根据文件类型提取文本"""
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith(('.docx', '.doc')):
            return self.extract_text_from_docx(file_path)
        else:
            return {"success": False, "error": "不支持的文件格式"}
    
    def smart_chunking(self, text: str, language: str, chunk_size: int = 500) -> List[str]:
        """根据语言智能分块"""
        if language == 'en':
            # 英文分块策略
            sentences = re.split(r'[.!?]+', text)
            chunks = []
            current_chunk = ""
            
            for sentence in sentences:
                clean_sentence = sentence.strip()
                if len(clean_sentence) == 0:
                    continue
                    
                if len(current_chunk) + len(clean_sentence) <= chunk_size:
                    current_chunk += clean_sentence + ". "
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = clean_sentence + ". "
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        else:
            # 中文分块策略
            paragraphs = re.split(r'\n\s*\n', text)
            chunks = []
            current_chunk = ""
            
            for paragraph in paragraphs:
                clean_para = re.sub(r'\s+', ' ', paragraph.strip())
                if len(clean_para) == 0:
                    continue
                    
                if len(current_chunk) + len(clean_para) <= chunk_size:
                    current_chunk += clean_para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = clean_para + "\n\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        
        return chunks
    
    def find_all_documents(self, root_path: str) -> List[Dict]:
        """递归查找所有PDF和Word文档"""
        documents = []
        
        for root, dirs, files in os.walk(root_path):
            # 跳过系统文件夹和临时文件夹
            dirs[:] = [d for d in dirs if not d.startswith('.') and not d.startswith('_')]
            
            for file in files:
                if file.lower().endswith(('.pdf', '.docx', '.doc')):
                    full_path = os.path.join(root, file)
                    # 计算相对路径用于分类
                    relative_path = os.path.relpath(root, root_path)
                    category = relative_path if relative_path != '.' else os.path.basename(root_path)
                    
                    documents.append({
                        'path': full_path,
                        'filename': file,
                        'category': category,
                        'folder_structure': relative_path
                    })
        
        return documents
    
    def process_document(self, doc_info: Dict) -> int:
        """处理单个文档并上传到云端"""
        file_path = doc_info['path']
        category = doc_info['category']
        filename = doc_info['filename']
        
        print(f"\n🔄 处理: {filename}")
        print(f"    📁 分类: {category}")
        
        # 提取文本
        result = self.extract_text(file_path)
        if not result["success"]:
            print(f"   ❌ 文本提取失败: {result.get('error', '未知错误')}")
            return 0
        
        text = result["text"]
        file_type = result["file_type"]
        
        if not text or len(text.strip()) < 50:
            print("   ⚠️  文档内容过少，跳过处理")
            return 0
        
        # 检测语言
        language = self.detect_language(text)
        
        # 根据语言分块
        chunks = self.smart_chunking(text, language)
        
        print(f"   📝 生成 {len(chunks)} 个文本块 | 语言: {language}")
        
        # 准备数据 - 分批上传以避免配额限制
        batch_size = 30  # 每批上传30个文档块
        total_uploaded = 0
        
        for batch_start in range(0, len(chunks), batch_size):
            batch_end = min(batch_start + batch_size, len(chunks))
            batch_chunks = chunks[batch_start:batch_end]
            
            documents = []
            metadatas = []
            ids = []
            
            for j, chunk in enumerate(batch_chunks):
                if len(chunk.strip()) > 30:
                    documents.append(chunk)
                    metadatas.append({
                        "category": category,
                        "doc_type": "医学文档",
                        "source_file": filename,
                        "file_type": file_type,
                        "language": language,
                        "chunk_index": batch_start + j,
                        "folder_path": doc_info['folder_structure'],
                        "full_path": file_path
                    })
                    # 使用短ID避免配额限制
                    short_id = self.generate_short_id(category, filename, language, batch_start + j)
                    ids.append(short_id)
            
            if documents:
                try:
                    # 关键修改：分批上传到云端
                    self.collection.add(
                        documents=documents,
                        metadatas=metadatas,
                        ids=ids
                    )
                    total_uploaded += len(documents)
                    print(f"   ✅ 批次 {batch_start//batch_size + 1}: 成功上传 {len(documents)} 个文本块")
                    
                    # 添加小延迟以避免速率限制
                    time.sleep(0.2)
                    
                except Exception as e:
                    print(f"   ❌ 批次上传失败: {e}")
                    # 继续处理下一批，不中断整个流程
        
        if total_uploaded > 0:
            print(f"   🎉 总计上传: {total_uploaded} 个文本块到云端！")
            return total_uploaded
        else:
            print("   ⚠️  没有有效的文本块可添加")
            return 0
    
    def get_collection_stats(self):
        """获取云端知识库统计信息（优化版，避免配额限制）"""
        try:
            count = self.collection.count()
            print(f"📊 数据库统计: 总文档数 = {count}")
            
            # 使用更小的limit避免配额问题
            sample_limit = min(50, count)  # 最多只取50条样本，避免超过300条限制
            language_dist = {}
            category_dist = {}
            
            if sample_limit > 0:
                try:
                    # 分批获取数据
                    sample_data = self.collection.peek(limit=sample_limit)
                    
                    for meta in sample_data.get("metadatas", []):
                        lang = meta.get("language", "unknown")
                        category = meta.get("category", "unknown")
                        
                        # 简化分类名称（去掉路径）
                        if isinstance(category, str) and '\\' in category:
                            # 只取最后一部分
                            category_parts = category.split('\\')
                            category = category_parts[-1]
                        
                        language_dist[lang] = language_dist.get(lang, 0) + 1
                        category_dist[category] = category_dist.get(category, 0) + 1
                        
                except Exception as e:
                    print(f"⚠️ 采样统计失败: {str(e)[:80]}")
                    # 如果采样失败，返回基本统计
                    return {
                        "total_chunks": count,
                        "language_distribution": {"unknown": count},
                        "category_distribution": {"unknown": count}
                    }
            else:
                # 空数据库
                language_dist = {}
                category_dist = {}
            
            return {
                "total_chunks": count,
                "language_distribution": language_dist,
                "category_distribution": category_dist
            }
            
        except Exception as e:
            print(f"❌ 获取统计失败: {str(e)[:100]}")
            # 返回安全的默认值
            return {
                "total_chunks": 0,
                "language_distribution": {},
                "category_distribution": {}
            }

class OptimizedFSHDUploader(FSHDBatchProcessor):
    """优化版上传器，添加进度保存和错误恢复"""
    
    def __init__(self, cloud_api_key: str, tenant_id: str, database_name: str = "FSHD"):
        super().__init__(cloud_api_key, tenant_id, database_name)
        self.progress_file = "upload_progress.json"
        self.load_progress()
    
    def load_progress(self):
        """加载上传进度"""
        try:
            if os.path.exists(self.progress_file):
                with open(self.progress_file, 'r', encoding='utf-8') as f:
                    self.progress = json.load(f)
                print(f"📚 加载进度: 已处理 {len(self.progress.get('processed_files', []))} 个文件")
            else:
                self.progress = {
                    "start_time": datetime.now().isoformat(),
                    "processed_files": [],
                    "failed_files": [],
                    "total_chunks": 0,
                    "last_checkpoint": None
                }
        except Exception as e:
            print(f"⚠️ 加载进度失败: {e}")
            self.progress = {"processed_files": [], "failed_files": []}
    
    def save_progress(self):
        """保存上传进度"""
        try:
            self.progress["last_checkpoint"] = datetime.now().isoformat()
            with open(self.progress_file, 'w', encoding='utf-8') as f:
                json.dump(self.progress, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"⚠️ 保存进度失败: {e}")
    
    def is_file_processed(self, file_path: str):
        """检查文件是否已处理"""
        return file_path in self.progress.get("processed_files", [])
    
    def process_entire_knowledge_base_safe(self, root_path: str):
        """安全处理整个知识库（支持断点续传）"""
        if not os.path.exists(root_path):
            print(f"❌ 知识库路径不存在: {root_path}")
            return
        
        print("🔍 扫描知识库文件夹结构...")
        all_documents = self.find_all_documents(root_path)
        
        if not all_documents:
            print("❌ 未找到任何PDF或Word文档")
            return
        
        print(f"\n📊 扫描完成！找到 {len(all_documents)} 个文档")
        
        # 过滤已处理的文件
        pending_documents = []
        for doc in all_documents:
            if not self.is_file_processed(doc['path']):
                pending_documents.append(doc)
        
        print(f"📋 待处理文件: {len(pending_documents)} 个 (跳过 {len(all_documents)-len(pending_documents)} 个已处理)")
        
        # 按分类统计
        category_stats = {}
        for doc in pending_documents:
            category = doc['category']
            category_stats[category] = category_stats.get(category, 0) + 1
        
        if category_stats:
            print("\n📂 待处理文档分类分布:")
            for category, count in sorted(category_stats.items()):
                print(f"   {category}: {count} 个文档")
        
        total_chunks = self.progress.get("total_chunks", 0)
        processed_files = len(self.progress.get("processed_files", []))
        
        print(f"\n{'='*60}")
        print("🚀 开始批量上传到云端知识库（安全模式）")
        print(f"📈 进度: {processed_files}/{len(all_documents)} 文件")
        print(f"{'='*60}")
        
        batch_counter = 0
        for i, doc_info in enumerate(pending_documents, 1):
            file_path = doc_info['path']
            
            print(f"\n[{processed_files + i}/{len(all_documents)}] ", end="")
            
            try:
                chunks_added = self.process_document(doc_info)
                
                if chunks_added > 0:
                    total_chunks += chunks_added
                    processed_files += 1
                    self.progress["processed_files"].append(file_path)
                    self.progress["total_chunks"] = total_chunks
                    
                    # 每处理5个文件保存一次进度
                    if i % 5 == 0:
                        self.save_progress()
                        print(f"   💾 进度已保存 ({processed_files}/{len(all_documents)})")
                else:
                    self.progress.setdefault("failed_files", []).append({
                        "path": file_path,
                        "reason": "无有效文本块",
                        "time": datetime.now().isoformat()
                    })
                    
            except Exception as e:
                error_msg = f"处理失败: {str(e)[:100]}"
                print(f"   ❌ {error_msg}")
                self.progress.setdefault("failed_files", []).append({
                    "path": file_path,
                    "reason": error_msg,
                    "time": datetime.now().isoformat()
                })
                # 继续处理下一个文件，不中断
                continue
            
            # 每处理10个文件稍作休息，避免速率限制
            batch_counter += 1
            if batch_counter >= 10:
                time.sleep(3)
                batch_counter = 0
        
        # 最终保存进度
        self.save_progress()
        
        print(f"\n{'🎉' * 30}")
        print("云端知识库批量上传完成！")
        print(f"{'🎉' * 30}")
        print(f"📊 处理统计:")
        print(f"   📁 扫描文档总数: {len(all_documents)} 个")
        print(f"   ✅ 成功处理: {processed_files} 个文档")
        print(f"   ❌ 失败: {len(self.progress.get('failed_files', []))} 个文档")
        print(f"   🧩 生成文本块: {total_chunks} 个")
        
        # 获取统计信息（使用修复后的方法）
        try:
            stats = self.get_collection_stats()
            print(f"\n📈 云端知识库总体统计:")
            print(f"   🧩 总文本块数: {stats['total_chunks']}")
            if stats.get('language_distribution'):
                print(f"   🌐 语言分布: {stats['language_distribution']}")
            if stats.get('category_distribution'):
                print(f"   📂 分类数量: {len(stats['category_distribution'])} 个")
        except Exception as e:
            print(f"\n⚠️  获取最终统计失败（不影响数据上传）: {str(e)[:80]}")
            print(f"   您可以通过查询API验证数据: curl http://localhost:5000/api/stats")
        
        # 如果有失败的文件
        failed_files = self.progress.get("failed_files", [])
        if failed_files:
            print(f"\n⚠️  失败文件列表 ({len(failed_files)} 个):")
            for fail in failed_files[:5]:  # 只显示前5个
                filename = os.path.basename(fail['path']) if 'path' in fail else '未知文件'
                print(f"   - {filename}: {fail.get('reason', '未知原因')}")
        
        # 保存详细报告
        report_file = f"upload_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        try:
            # 尝试获取最终统计用于报告
            try:
                final_stats = self.get_collection_stats()
            except:
                final_stats = {"total_chunks": "获取失败", "error": "配额限制"}
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "summary": {
                        "total_documents": len(all_documents),
                        "processed_success": processed_files,
                        "failed": len(failed_files),
                        "total_chunks": total_chunks,
                        "database_count": final_stats.get('total_chunks', '未知')
                    },
                    "progress": self.progress,
                    "stats": final_stats
                }, f, ensure_ascii=False, indent=2)
            print(f"\n📄 详细报告已保存: {report_file}")
        except Exception as e:
            print(f"\n⚠️  保存报告失败: {e}")

def main_optimized():
    """优化版主函数"""
    # 🎯 使用您的云端凭据
    CLOUD_API_KEY = "ck-G1qMBnQAHG1B1xZN8b1fHzcjbq1TxdbSFsNofzGaZT5c"
    TENANT_ID = "bf4422ea-4e6b-4f9b-8682-bc9f92d22f04"
    
    # 初始化优化版上传器
    uploader = OptimizedFSHDUploader(
        cloud_api_key=CLOUD_API_KEY,
        tenant_id=TENANT_ID,
        database_name="FSHD"
    )
    
    # 知识库根路径
    knowledge_base_path = r"C:\yoyo\openrd-master\FSHD_知识库"
    
    print("🎯 开始安全批量上传FSHD知识库到云端")
    print(f"📍 本地知识库位置: {knowledge_base_path}")
    print("🛡️  模式: 支持断点续传、错误恢复、进度保存")
    print("⏳ 上传可能需要一些时间，请耐心等待...\n")
    
    # 安全处理整个知识库
    uploader.process_entire_knowledge_base_safe(knowledge_base_path)

if __name__ == "__main__":
    main_optimized()