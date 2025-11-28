import chromadb
import os
import PyPDF2
from docx import Document
from typing import List, Dict
import re
import langdetect
from langdetect import detect, DetectorFactory

DetectorFactory.seed = 0

class FSHDBatchProcessor:
    def __init__(self):
        self.client = chromadb.PersistentClient(path="./chroma_db")
        self.collection = self.client.get_or_create_collection("fshd_knowledge_base")
        print("🚀 FSHD批量知识库处理器初始化完成！")
    
    def detect_language(self, text: str) -> str:
        """检测文本语言"""
        try:
            sample_text = text[:1000] if len(text) > 1000 else text
            if len(sample_text.strip()) < 10:  
                return "unknown"
            return detect(sample_text)
        except:
            return "unknown"
    
    def extract_text_from_pdf(self, pdf_path: str) -> Dict:
        """从PDF提取文本"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                total_pages = len(reader.pages)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
            return {
                "text": text,
                "success": True,
                "pages": total_pages,
                "file_type": "pdf"
            }
        except Exception as e:
            print(f"❌ 读取PDF失败 {pdf_path}: {e}")
            return {"success": False, "error": str(e)}
    
    def extract_text_from_docx(self, docx_path: str) -> Dict:
        """从Word文档提取文本"""
        text = ""
        try:
            doc = Document(docx_path)
            total_paragraphs = len(doc.paragraphs)
            for para in doc.paragraphs:
                text += para.text + "\n"
            return {
                "text": text,
                "success": True,
                "paragraphs": total_paragraphs,
                "file_type": "docx"
            }
        except Exception as e:
            print(f"❌ 读取Word文档失败 {docx_path}: {e}")
            return {"success": False, "error": str(e)}
    
    def extract_text(self, file_path: str) -> Dict:
        """根据文件类型提取文本"""
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith(('.docx', '.doc')):
            return self.extract_text_from_docx(file_path)
        else:
            return {"success": False, "error": "不支持的文件格式"}
    
    def smart_chunking(self, text: str, language: str, chunk_size: int = 500) -> List[str]:
        """根据语言智能分块"""
        if language == 'en':
            # 英文分块策略
            sentences = re.split(r'[.!?]+', text)
            chunks = []
            current_chunk = ""
            
            for sentence in sentences:
                clean_sentence = sentence.strip()
                if len(clean_sentence) == 0:
                    continue
                    
                if len(current_chunk) + len(clean_sentence) <= chunk_size:
                    current_chunk += clean_sentence + ". "
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = clean_sentence + ". "
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        else:
            # 中文分块策略
            paragraphs = re.split(r'\n\s*\n', text)
            chunks = []
            current_chunk = ""
            
            for paragraph in paragraphs:
                clean_para = re.sub(r'\s+', ' ', paragraph.strip())
                if len(clean_para) == 0:
                    continue
                    
                if len(current_chunk) + len(clean_para) <= chunk_size:
                    current_chunk += clean_para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = clean_para + "\n\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        
        return chunks
    
    def find_all_documents(self, root_path: str) -> List[Dict]:
        """递归查找所有PDF和Word文档"""
        documents = []
        
        for root, dirs, files in os.walk(root_path):
            # 跳过系统文件夹和临时文件夹
            dirs[:] = [d for d in dirs if not d.startswith('.') and not d.startswith('_')]
            
            for file in files:
                if file.lower().endswith(('.pdf', '.docx', '.doc')):
                    full_path = os.path.join(root, file)
                    # 计算相对路径用于分类
                    relative_path = os.path.relpath(root, root_path)
                    category = relative_path if relative_path != '.' else os.path.basename(root_path)
                    
                    documents.append({
                        'path': full_path,
                        'filename': file,
                        'category': category,
                        'folder_structure': relative_path
                    })
        
        return documents
    
    def process_document(self, doc_info: Dict) -> int:
        """处理单个文档"""
        file_path = doc_info['path']
        category = doc_info['category']
        filename = doc_info['filename']
        
        print(f"\n🔄 处理: {filename}")
        print(f"    📁 分类: {category}")
        print(f"    📂 路径: {file_path}")
        
        # 提取文本
        result = self.extract_text(file_path)
        if not result["success"]:
            print(f"   ❌ 文本提取失败: {result.get('error', '未知错误')}")
            return 0
        
        text = result["text"]
        file_type = result["file_type"]
        
        if not text or len(text.strip()) < 50:
            print("   ⚠️  文档内容过少，跳过处理")
            return 0
        
        # 检测语言
        language = self.detect_language(text)
        
        # 根据语言分块
        chunks = self.smart_chunking(text, language)
        
        # 添加文件特定信息
        file_info = ""
        if file_type == "pdf":
            file_info = f"页数: {result['pages']}"
        elif file_type == "docx":
            file_info = f"段落: {result['paragraphs']}"
        
        print(f"   📝 生成 {len(chunks)} 个文本块 | 语言: {language} | {file_info}")
        
        # 准备数据
        documents = []
        metadatas = []
        ids = []
        
        for j, chunk in enumerate(chunks):
            if len(chunk.strip()) > 30:
                documents.append(chunk)
                metadatas.append({
                    "category": category,
                    "doc_type": "医学文档",
                    "source_file": filename,
                    "file_type": file_type,
                    "language": language,
                    "chunk_index": j,
                    "folder_path": doc_info['folder_structure'],
                    "full_path": file_path
                })
                ids.append(f"{category}_{filename}_{language}_{j}")
        
        if documents:
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            print(f"   ✅ 成功添加: {len(documents)} 个文本块")
            return len(documents)
        else:
            print("   ⚠️  没有有效的文本块可添加")
            return 0
    
    def process_entire_knowledge_base(self, root_path: str):
        """处理整个知识库文件夹"""
        if not os.path.exists(root_path):
            print(f"❌ 知识库路径不存在: {root_path}")
            return
        
        print("🔍 扫描知识库文件夹结构...")
        all_documents = self.find_all_documents(root_path)
        
        if not all_documents:
            print("❌ 未找到任何PDF或Word文档")
            return
        
        print(f"\n📊 扫描完成！找到 {len(all_documents)} 个文档")
        
        # 按分类统计
        category_stats = {}
        for doc in all_documents:
            category = doc['category']
            category_stats[category] = category_stats.get(category, 0) + 1
        
        print("\n📂 文档分类分布:")
        for category, count in category_stats.items():
            print(f"   {category}: {count} 个文档")
        
        total_chunks = 0
        processed_files = 0
        
        print(f"\n{'='*60}")
        print("🚀 开始批量处理文档...")
        print(f"{'='*60}")
        
        for i, doc_info in enumerate(all_documents, 1):
            print(f"\n[{i}/{len(all_documents)}] ", end="")
            chunks_added = self.process_document(doc_info)
            total_chunks += chunks_added
            if chunks_added > 0:
                processed_files += 1
        
        # 显示最终统计
        stats = self.get_collection_stats()
        
        print(f"\n{'🎉' * 30}")
        print("批量处理完成！")
        print(f"{'🎉' * 30}")
        print(f"📊 处理统计:")
        print(f"   📁 扫描文档总数: {len(all_documents)} 个")
        print(f"   ✅ 成功处理: {processed_files} 个文档")
        print(f"   🧩 生成文本块: {total_chunks} 个")
        print(f"\n📈 知识库总体统计:")
        print(f"   🧩 总文本块数: {stats['total_chunks']}")
        print(f"   🌐 语言分布: {stats['language_distribution']}")
        print(f"   📂 分类数量: {len(stats['category_distribution'])} 个")
        print(f"   📋 分类详情: {stats['category_distribution']}")
    
    def search_fshd_knowledge(self, question: str, n_results: int = 3, language_filter: str = None):
        """搜索FSHD知识库"""
        where_filter = None
        if language_filter:
            where_filter = {"language": language_filter}
        
        results = self.collection.query(
            query_texts=[question],
            n_results=n_results,
            where=where_filter
        )
        return results
    
    def search_knowledge(self, question: str, n_results: int = 3, language_filter: str = None):
        """搜索FSHD知识库（兼容性方法）"""
        return self.search_fshd_knowledge(question, n_results, language_filter)
    
    def get_collection_stats(self):
        """获取知识库统计信息"""
        count = self.collection.count()
        all_metadatas = self.collection.get()["metadatas"]
        
        language_dist = {}
        category_dist = {}
        file_type_dist = {}
        
        for meta in all_metadatas:
            lang = meta.get("language", "unknown")
            category = meta.get("category", "unknown")
            file_type = meta.get("file_type", "unknown")
            
            language_dist[lang] = language_dist.get(lang, 0) + 1
            category_dist[category] = category_dist.get(category, 0) + 1
            file_type_dist[file_type] = file_type_dist.get(file_type, 0) + 1
        
        return {
            "total_chunks": count,
            "language_distribution": language_dist,
            "category_distribution": category_dist,
            "file_type_distribution": file_type_dist
        }

def main():
    """主函数 - 批量处理整个FSHD知识库"""
    processor = FSHDBatchProcessor()
    
    # 知识库根路径
    knowledge_base_path = r"C:\yoyo\openrd-master\FSHD_知识库"
    
    print("🎯 开始批量处理FSHD知识库")
    print(f"📍 知识库位置: {knowledge_base_path}")
    print("⏳ 这可能需要一些时间，请耐心等待...\n")
    
    # 批量处理整个知识库
    processor.process_entire_knowledge_base(knowledge_base_path)
    
    # 测试搜索
    print(f"\n🔍 测试搜索功能...")
    test_questions = [
        "What is Facioscapulohumeral Muscular Dystrophy?",
        "FSHD的诊断方法有哪些？",
        "FSHD genetic testing"
    ]
    
    for question in test_questions:
        results = processor.search_fshd_knowledge(question, n_results=2)
        print(f"\n❓ 问题: {question}")
        print(f"📋 找到 {len(results['documents'][0])} 个相关结果")
        for j, doc in enumerate(results['documents'][0]):
            print(f"   {j+1}. {doc[:150]}...")
            print(f"      分类: {results['metadatas'][0][j].get('category', 'unknown')}")
            print(f"      语言: {results['metadatas'][0][j].get('language', 'unknown')}")
            print(f"      来源: {results['metadatas'][0][j].get('source_file', 'unknown')}")

if __name__ == "__main__":
    main()