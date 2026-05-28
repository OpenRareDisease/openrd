"""Microsoft Word (.docx) parser via python-docx, with a zip+XML
fallback for files that confuse python-docx (broken bookmarks etc.).

Strategy: walk the document body once, group paragraphs by the most
recent `Heading 1` / `Heading 2` they sit under, and emit one
`ParsedSection` per heading group (labelled with the heading text).
Documents with no headings degenerate to a single section that holds
the whole body — the chunker handles further splitting.

Tables are flattened to tab-separated text and appended to the
section they belong to. Inline images are ignored on purpose; the
image parser handles standalone images and a follow-up could extract
embedded images if a clinical use case appears.

Two failure modes we handle explicitly:
  - Legacy .doc files saved with a .docx extension: detected by the
    Composite Document File magic (`D0 CF 11 E0`); we surface a clear
    parse_error so the operator can re-save as real .docx rather than
    silently losing the file.
  - .docx whose bookmark XML refs python-docx can't resolve: we fall
    back to opening the file as a zip and stripping XML tags from
    `word/document.xml`. Loses heading structure but recovers text,
    which is what the embedder cares about.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import List

from .base import Parser, ParseResult, ParsedSection

#: Magic bytes for a legacy .doc (OLE2 compound file). Files with this
#: header are not valid .docx and python-docx will raise PackageNotFound.
_OLE2_MAGIC = b"\xD0\xCF\x11\xE0"

#: Lazy strip of XML tags. Good enough for fallback text recovery —
#: we lose paragraph/run boundaries but keep the words.
_XML_TAG_RE = re.compile(r"<[^>]+>")


class DocxParser(Parser):
    extensions = (".docx",)
    parser_name = "docx"

    def parse(self, path: Path) -> ParseResult:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return ParseResult(
                sections=[],
                metadata={
                    "parser": self.parser_name,
                    "parse_error": "python-docx not installed",
                },
            )

        # Detect legacy .doc files masquerading as .docx before
        # python-docx prints a less actionable error.
        try:
            # `with` so the file descriptor is closed when the read
            # finishes — the previous `path.open("rb").read(4)` left
            # the fd open until CPython's GC happened to collect it.
            # Harmless on CPython under normal load, but on PyPy or
            # with low fd ulimits a corpus-wide ingest could exhaust
            # the limit and start raising OSError partway through.
            with path.open("rb") as fh:
                head = fh.read(4)
        except Exception as exc:
            return ParseResult(
                sections=[],
                metadata={
                    "parser": self.parser_name,
                    "parse_error": f"cannot read file: {exc}",
                },
            )
        if head == _OLE2_MAGIC:
            return ParseResult(
                sections=[],
                metadata={
                    "parser": self.parser_name,
                    "parse_error": (
                        "legacy .doc format with .docx extension; "
                        "re-save as .docx in Word/WPS before re-ingesting"
                    ),
                    "format_hint": "legacy_doc",
                },
            )

        try:
            doc = Document(str(path))
        except Exception as exc:
            # Try the zip+XML fallback. Catches broken bookmark refs,
            # missing item-name entries etc. — anything where the file
            # is still a valid zip but python-docx's relationship
            # graph can't be reconstructed.
            fallback = _parse_via_xml(path)
            if fallback is not None:
                fallback.metadata["fallback_reason"] = str(exc)
                return fallback
            return ParseResult(
                sections=[],
                metadata={
                    "parser": self.parser_name,
                    "parse_error": f"python-docx failed to open: {exc}",
                },
            )

        sections: List[ParsedSection] = []
        current_heading = ""
        current_lines: List[str] = []

        def flush() -> None:
            text = "\n".join(line for line in current_lines if line.strip()).strip()
            if text:
                sections.append(
                    ParsedSection(
                        text=text,
                        label=current_heading,
                        extra={"section_index": len(sections)},
                    )
                )

        # Walk the document body in document order so a table that
        # appears between Heading A and Heading B is attached to A,
        # not to whatever heading is current after every paragraph
        # has been processed. python-docx's `doc.paragraphs` /
        # `doc.tables` collections each preserve internal order but
        # lose ordering between the two types; `iter_inner_content`
        # is the documented escape hatch (>= 1.0).
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore

        for block in doc.iter_inner_content():
            if isinstance(block, Paragraph):
                style_name = (block.style.name if block.style else "") or ""
                text = (block.text or "").strip()
                if not text:
                    continue
                if style_name.startswith("Heading"):
                    flush()
                    current_heading = text
                    current_lines = []
                    continue
                current_lines.append(text)
            elif isinstance(block, Table):
                rows: List[str] = []
                for row in block.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        rows.append("\t".join(cells))
                if rows:
                    current_lines.append("[表格]")
                    current_lines.extend(rows)
            # Other inner-content types (sdt, etc.) are ignored.

        flush()

        return ParseResult(
            sections=sections,
            metadata={
                "parser": self.parser_name,
                "sections_total": len(sections),
            },
        )


def _parse_via_xml(path: Path) -> ParseResult | None:
    """Last-resort text extraction: open the .docx as a zip, read
    `word/document.xml`, strip XML tags. Returns None when even the
    zip can't be opened (so the caller surfaces the original
    python-docx error instead of a misleading fallback message).
    Section boundaries are lost — everything becomes a single
    ParsedSection that the chunker will paragraph-split.
    """
    try:
        with zipfile.ZipFile(path) as zf:
            try:
                raw = zf.read("word/document.xml")
            except KeyError:
                return None
    except Exception:
        return None

    # Decode + drop tags. The XML uses unicode space tokens we want to
    # preserve, but tag attributes leak garbage if we just strip
    # angle brackets — handle the common case of paragraph close.
    text = raw.decode("utf-8", errors="ignore")
    # Insert newline at paragraph and table boundaries so we keep at
    # least some structure for the chunker.
    text = re.sub(r"</w:p>", "\n", text)
    text = re.sub(r"</w:tr>", "\n", text)
    text = _XML_TAG_RE.sub("", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if not text:
        return None

    return ParseResult(
        sections=[
            ParsedSection(
                text=text,
                label="",
                extra={"source_method": "xml_fallback"},
            )
        ],
        metadata={
            "parser": "docx",
            "fallback": "xml",
        },
    )
