"""Tests for the docx parser.

Builds tiny .docx fixtures on the fly via python-docx so we don't
have to ship binary blobs in the repo. Covers heading-aware
sectioning, table flattening, and the legacy-.doc detection.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document  # type: ignore

from kb_parsers.docx_parser import DocxParser


def _write_docx(tmp_path: Path, build) -> Path:
    f = tmp_path / "fixture.docx"
    doc = Document()
    build(doc)
    doc.save(str(f))
    return f


def test_docx_parser_groups_paragraphs_by_heading(tmp_path: Path) -> None:
    def build(doc):
        doc.add_heading("First section", level=1)
        doc.add_paragraph("alpha line")
        doc.add_paragraph("beta line")
        doc.add_heading("Second section", level=1)
        doc.add_paragraph("gamma line")

    f = _write_docx(tmp_path, build)
    result = DocxParser().parse(f)
    labels = [s.label for s in result.sections]
    assert "First section" in labels
    assert "Second section" in labels

    first = next(s for s in result.sections if s.label == "First section")
    assert "alpha line" in first.text
    assert "beta line" in first.text
    assert "gamma line" not in first.text


def test_docx_parser_no_headings_yields_single_section(tmp_path: Path) -> None:
    def build(doc):
        doc.add_paragraph("plain one")
        doc.add_paragraph("plain two")

    f = _write_docx(tmp_path, build)
    result = DocxParser().parse(f)
    assert len(result.sections) == 1
    assert result.sections[0].label == ""
    assert "plain one" in result.sections[0].text
    assert "plain two" in result.sections[0].text


def test_docx_parser_flattens_tables(tmp_path: Path) -> None:
    def build(doc):
        doc.add_paragraph("intro paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "h1"
        table.cell(0, 1).text = "h2"
        table.cell(1, 0).text = "v1"
        table.cell(1, 1).text = "v2"

    f = _write_docx(tmp_path, build)
    result = DocxParser().parse(f)
    body = "\n".join(s.text for s in result.sections)
    assert "h1\th2" in body
    assert "v1\tv2" in body
    assert "[表格]" in body


def test_docx_parser_skips_empty_paragraphs(tmp_path: Path) -> None:
    def build(doc):
        doc.add_paragraph("real content")
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("more")

    f = _write_docx(tmp_path, build)
    result = DocxParser().parse(f)
    assert result.sections
    text = result.sections[0].text
    assert "real content" in text
    assert "more" in text


def test_docx_parser_detects_legacy_doc_format(tmp_path: Path) -> None:
    # Forge a file with OLE2 magic bytes that has a .docx extension —
    # this is the common WPS / older-Word failure mode.
    f = tmp_path / "looks_legit.docx"
    f.write_bytes(b"\xD0\xCF\x11\xE0" + b"\x00" * 100)
    result = DocxParser().parse(f)
    assert result.sections == []
    assert result.metadata.get("format_hint") == "legacy_doc"
    assert "legacy .doc" in result.metadata.get("parse_error", "")


def test_docx_parser_corrupt_file_returns_parse_error(tmp_path: Path) -> None:
    f = tmp_path / "broken.docx"
    f.write_bytes(b"not actually a zip")
    result = DocxParser().parse(f)
    assert result.sections == []
    assert "parse_error" in result.metadata


def test_docx_parser_attaches_tables_to_the_current_heading(tmp_path: Path) -> None:
    """Regression: tables used to be iterated AFTER all paragraphs,
    so every table in the document ended up attached to whatever
    heading was last seen during the paragraph pass. A table that
    appears between Heading A and Heading B should belong to A."""

    def build(doc):
        doc.add_heading("Heading A", level=1)
        doc.add_paragraph("alpha intro")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "alpha_h1"
        table.cell(0, 1).text = "alpha_h2"
        table.cell(1, 0).text = "alpha_v1"
        table.cell(1, 1).text = "alpha_v2"
        doc.add_heading("Heading B", level=1)
        doc.add_paragraph("beta paragraph")

    f = _write_docx(tmp_path, build)
    result = DocxParser().parse(f)

    by_label = {s.label: s.text for s in result.sections}
    assert "Heading A" in by_label, f"section for A missing: {list(by_label)}"
    assert "Heading B" in by_label, f"section for B missing: {list(by_label)}"

    # Table content must live under A (where it appeared), not B.
    assert "alpha_h1\talpha_h2" in by_label["Heading A"]
    assert "alpha_v1\talpha_v2" in by_label["Heading A"]
    assert "alpha_h1" not in by_label["Heading B"]
    assert "beta paragraph" in by_label["Heading B"]
    assert "beta paragraph" not in by_label["Heading A"]


def test_docx_parser_handles_interleaved_paragraphs_and_tables(tmp_path: Path) -> None:
    """Multiple tables under different headings should each land in
    their own section. Exercises the document-order traversal end
    to end."""

    def build(doc):
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("text 1")
        t1 = doc.add_table(rows=1, cols=1)
        t1.cell(0, 0).text = "t1_cell"
        doc.add_heading("Section 2", level=1)
        t2 = doc.add_table(rows=1, cols=1)
        t2.cell(0, 0).text = "t2_cell"
        doc.add_paragraph("text 2")

    f = _write_docx(tmp_path, build)
    result = DocxParser().parse(f)
    by_label = {s.label: s.text for s in result.sections}

    assert "t1_cell" in by_label["Section 1"]
    assert "t2_cell" in by_label["Section 2"]
    assert "t1_cell" not in by_label["Section 2"]
    assert "t2_cell" not in by_label["Section 1"]
    assert "text 1" in by_label["Section 1"]
    assert "text 2" in by_label["Section 2"]
